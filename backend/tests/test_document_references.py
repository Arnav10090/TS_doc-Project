from app.generation.document_references import collect_document_references
from app.generation.document_references import apply_document_references
from docx import Document
from docx.oxml import OxmlElement
from zipfile import ZipFile


def test_collect_document_references_numbers_builtin_tables_and_custom_items(tmp_path):
    project_id = "project-1"
    sections = {
        "revision_history": {"rows": []},
        "executive_summary": {},
        "system_config": {},
        "custom_section_1700000000000_12345678-1234-1234-1234-123456789abc": {
            "title": "Custom Section",
            "insertAfterKey": "system_config",
            "subsections": [
                {
                    "key": "custom_subsection_1",
                    "name": "Network Layout",
                    "contentType": "image",
                    "data": {
                        "images": [
                            {
                                "base64": "data:image/png;base64,abcd",
                                "filename": "network.png",
                                "mimeType": "image/png",
                                "caption": "Network Layout",
                            }
                        ]
                    },
                },
                {
                    "key": "custom_subsection_2",
                    "name": "Training Schedule",
                    "contentType": "table",
                    "data": {
                        "tables": [
                            {
                                "caption": "Training Schedule",
                                "columns": ["Day", "Topic"],
                                "rows": [{"Day": "1", "Topic": "Overview"}],
                            }
                        ]
                    },
                },
            ],
        },
    }
    image_dir = tmp_path / "images" / project_id
    image_dir.mkdir(parents=True)
    (image_dir / "architecture.png").write_bytes(b"fake image")

    references = collect_document_references(sections, str(tmp_path), project_id)

    assert [(figure.number, figure.name) for figure in references.figures] == [
        (1, "System Architecture"),
        (2, "Network Layout"),
    ]
    assert [(table.number, table.name) for table in references.tables] == [
        (1, "Revision History"),
        (2, "Training Schedule"),
    ]


def test_collect_document_references_renumbers_later_builtin_tables(tmp_path):
    references = collect_document_references(
        {
            "abbreviations": {"rows": []},
            "tech_stack": {"rows": []},
            "custom_section_1700000000000_12345678-1234-1234-1234-123456789abc": {
                "title": "",
                "insertAfterKey": "abbreviations",
                "displayMode": "subsection",
                "subsections": [
                    {
                        "key": "custom_subsection_1",
                        "name": "Inserted Data",
                        "contentType": "table",
                        "data": {
                            "tables": [
                                {
                                    "caption": "Inserted Data",
                                    "columns": ["Name"],
                                    "rows": [{"Name": "Example"}],
                                }
                            ]
                        },
                    }
                ],
            },
        },
        str(tmp_path),
        "project-1",
    )

    assert [(table.number, table.name) for table in references.tables] == [
        (1, "Abbreviations Used"),
        (2, "Inserted Data"),
        (3, "Technology Stack"),
    ]


def test_apply_document_references_adds_captions_and_final_lists(tmp_path):
    doc = Document()
    for index in range(10):
        doc.add_paragraph(f"Section {index + 1}", style="Heading 1")

    for index in range(8):
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = f"Table {index}"

    file_path = tmp_path / "document.docx"
    doc.save(file_path)

    apply_document_references(
        str(file_path),
        {
            "revision_history": {"rows": []},
            "executive_summary": {},
            "abbreviations": {"rows": []},
        },
        str(tmp_path),
        "project-1",
    )

    result = Document(file_path)
    text = "\n".join(paragraph.text for paragraph in result.paragraphs)

    assert "Table 1: Revision History" in text
    assert "Table 2: Abbreviations Used" in text
    assert "Client Reference Logos" not in text
    assert "11. List of Figures and Tables" in text
    assert "11.1 List of Figures" in text
    assert "11.2 List of Tables" in text


def test_apply_document_references_adds_dynamic_toc_field(tmp_path):
    doc = Document()
    doc.add_paragraph("TABLE OF CONTENTS")
    doc.add_paragraph("EXECUTIVE SUMMARY", style="Heading 1")

    file_path = tmp_path / "document.docx"
    doc.save(file_path)

    apply_document_references(str(file_path), {}, str(tmp_path), "project-1")

    with ZipFile(file_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")
        settings_xml = docx.read("word/settings.xml").decode("utf-8")

    assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml
    assert "w:updateFields" in settings_xml


def test_apply_document_references_handles_existing_toc_field(tmp_path):
    doc = Document()
    toc_heading = doc.add_paragraph("TABLE OF CONTENTS")

    field_paragraph_element = OxmlElement("w:p")
    toc_heading._p.addnext(field_paragraph_element)

    run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    run.append(instruction)
    field_paragraph_element.append(run)

    doc.add_paragraph("EXECUTIVE SUMMARY", style="Heading 1")

    file_path = tmp_path / "document.docx"
    doc.save(file_path)

    apply_document_references(str(file_path), {}, str(tmp_path), "project-1")

    with ZipFile(file_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")
        settings_xml = docx.read("word/settings.xml").decode("utf-8")

    assert document_xml.count('TOC \\o "1-3" \\h \\z \\u') == 1
    assert "w:updateFields" in settings_xml


def test_apply_document_references_inserts_custom_content_at_preview_anchors(tmp_path):
    doc = Document()
    doc.add_paragraph("EXECUTIVE SUMMARY", style="Heading 1")
    doc.add_paragraph("GENERAL OVERVIEW", style="Heading 1")
    doc.add_paragraph("INTRODUCTION", style="Heading 2")
    doc.add_paragraph("Introduction body")
    doc.add_paragraph("ABBREVIATIONS USED", style="Heading 2")
    doc.add_paragraph("Abbreviations body")
    doc.add_paragraph("TECHNOLOGY STACK", style="Heading 1")
    doc.add_paragraph("THIRD PARTY SOFTWARE REQUIREMENTS", style="Heading 3")
    doc.add_paragraph("Third party software body")
    doc.add_paragraph("SCHEDULE", style="Heading 1")

    file_path = tmp_path / "document.docx"
    doc.save(file_path)

    apply_document_references(
        str(file_path),
        {
            "introduction": {},
            "abbreviations": {},
            "third_party_sw": {},
            "custom_section_1700000000000_12345678-1234-1234-1234-123456789abc": {
                "title": "",
                "displayMode": "subsection",
                "insertAfterKey": "introduction",
                "subsections": [
                    {
                        "key": "custom_subsection_1",
                        "name": "Inserted Intro Detail",
                        "contentType": "paragraph",
                        "data": {
                            "paragraphs": [
                                {"html": "<ul><li><p>First bullet</p></li></ul>"}
                            ]
                        },
                    }
                ],
            },
            "custom_section_1700000001000_22345678-1234-1234-1234-123456789abc": {
                "title": "Section 5",
                "displayMode": "section",
                "insertAfterKey": "third_party_sw",
                "subsections": [
                    {
                        "key": "custom_subsection_2",
                        "name": "Inserted Section Detail",
                        "contentType": "paragraph",
                        "data": {"paragraphs": [{"html": "<p>Section body</p>"}]},
                    }
                ],
            },
        },
        str(tmp_path),
        "project-1",
    )

    result = Document(file_path)
    paragraphs = [paragraph for paragraph in result.paragraphs if paragraph.text.strip()]
    text = [paragraph.text.strip() for paragraph in paragraphs]

    assert text.index("INTRODUCTION") < text.index("Inserted Intro Detail")
    assert text.index("Inserted Intro Detail") < text.index("ABBREVIATIONS USED")
    assert text.index("THIRD PARTY SOFTWARE REQUIREMENTS") < text.index("Section 5")
    assert text.index("Section 5") < text.index("SCHEDULE")
    assert not any(item.startswith("Custom Section") for item in text)

    inserted_intro = paragraphs[text.index("Inserted Intro Detail")]
    inserted_section = paragraphs[text.index("Section 5")]
    first_bullet = paragraphs[text.index("First bullet")]
    assert inserted_intro.style.name == "Heading 2"
    assert inserted_section.style.name == "Heading 1"
    assert first_bullet.style.name in {"List Bullet", "Normal"}
