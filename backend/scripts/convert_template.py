"""
Word Template Conversion Script for TS Document Generator

This script converts the original Word template (TS_Template_original.docx) 
into a Jinja2-compatible template (TS_Template_jinja.docx) by replacing 
static text placeholders with Jinja2 template variables.

CRITICAL: Variable names are case-sensitive. Uppercase/lowercase mismatches
will cause content to disappear with no error during document generation.
"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import copy

# Define input and output paths
INPUT = Path("templates/TS_Template_original.docx")
OUTPUT = Path("templates/TS_Template_jinja.docx")

# Define all text replacements (80+ mappings)
REPLACEMENTS = {
    # Project metadata
    "SolutionFullName": "{{ SolutionFullName }}",
    "SolutionName": "{{ SolutionName }}",
    "SolutionAbbreviation": "{{ SolutionAbbreviation }}",
    "CLIENTNAME": "{{ ClientName }}",
    "ClientName": "{{ ClientName }}",
    "ClientLocation": "{{ ClientLocation }}",
    "CLIENTLOCATION": "{{ ClientLocation }}",
    "ClientAbbreviation": "{{ ClientAbbreviation }}",
    "RefNumber": "{{ RefNumber }}",
    "DocDate": "{{ DocDate }}",
    "DocVersion": "{{ DocVersion }}",
    
    # Executive Summary
    "ExecutiveSummaryPara1": "{{ ExecutiveSummaryPara1 }}",
    
    # Introduction
    "TenderReference": "{{ TenderReference }}",
    "TenderDate": "{{ TenderDate }}",
    
    # Process Flow
    "ProcessFlowDescription": "{{ ProcessFlowDescription }}",
    
    # Overview section
    "SystemObjective": "{{ SystemObjective }}",
    "ExistingSystemDescription": "{{ ExistingSystemDescription }}",
    "IntegrationDescription": "{{ IntegrationDescription }}",
    "TangibleBenefits": "{{ TangibleBenefits }}",
    "IntangibleBenefits": "{{ IntangibleBenefits }}",
    
    # Features (6 features with title, brief, description)
    "Feature1Title": "{{ Feature1Title }}",
    "Feature1Brief": "{{ Feature1Brief }}",
    "Feature1Description": "{{ Feature1Description }}",
    "Feature2Title": "{{ Feature2Title }}",
    "Feature2Brief": "{{ Feature2Brief }}",
    "Feature2Description": "{{ Feature2Description }}",
    "Feature3Title": "{{ Feature3Title }}",
    "Feature3Brief": "{{ Feature3Brief }}",
    "Feature3Description": "{{ Feature3Description }}",
    "Feature4Title": "{{ Feature4Title }}",
    "Feature4Brief": "{{ Feature4Brief }}",
    "Feature4Description": "{{ Feature4Description }}",
    "Feature5Title": "{{ Feature5Title }}",
    "Feature5Brief": "{{ Feature5Brief }}",
    "Feature5Description": "{{ Feature5Description }}",
    "Feature6Title": "{{ Feature6Title }}",
    "Feature6Brief": "{{ Feature6Brief }}",
    "Feature6Description": "{{ Feature6Description }}",
    
    # Remote Support - full paragraph replacement
    "RemoteSupportText": "{{ RemoteSupportText }}",
    
    # Customer Training
    "TrainingPersons": "{{ TrainingPersons }}",
    "TrainingDays": "{{ TrainingDays }}",
    
    # FAT Condition
    "FATCondition": "{{ FATCondition }}",
    
    # CRITICAL: Image placeholders (must be lowercase in Jinja2 variables)
    "ARCHITECTURE_DIAGRAM": "{{ architecture_diagram }}",
    "OVERALL_GANTT_CHART": "{{ overall_gantt }}",
    "SHUTDOWN_GANTT_CHART": "{{ shutdown_gantt }}",
    
    # Technology Stack (6 rows)
    "TS1_Component": "{{ TS1_Component }}",
    "TS1_Technology": "{{ TS1_Technology }}",
    "TS2_Component": "{{ TS2_Component }}",
    "TS2_Technology": "{{ TS2_Technology }}",
    "TS3_Component": "{{ TS3_Component }}",
    "TS3_Technology": "{{ TS3_Technology }}",
    "TS4_Component": "{{ TS4_Component }}",
    "TS4_Technology": "{{ TS4_Technology }}",
    "TS5_Component": "{{ TS5_Component }}",
    "TS5_Technology": "{{ TS5_Technology }}",
    "TS6_Component": "{{ TS6_Component }}",
    "TS6_Technology": "{{ TS6_Technology }}",
    
    # Hardware Specs Row 1 (4 spec lines + maker + qty)
    "HW1_Specs_Line1": "{{ HW1_Specs_Line1 }}",
    "HW1_Specs_Line2": "{{ HW1_Specs_Line2 }}",
    "HW1_Specs_Line3": "{{ HW1_Specs_Line3 }}",
    "HW1_Specs_Line4": "{{ HW1_Specs_Line4 }}",
    "HW1_Maker": "{{ HW1_Maker }}",
    "HW1_Qty": "{{ HW1_Qty }}",
    
    # Hardware Specs Rows 2-6 (varying spec line counts)
    "HW2_Specs": "{{ HW2_Specs }}",
    "HW2_Maker": "{{ HW2_Maker }}",
    "HW2_Qty": "{{ HW2_Qty }}",
    
    "HW3_Specs": "{{ HW3_Specs }}",
    "HW3_Maker": "{{ HW3_Maker }}",
    "HW3_Qty": "{{ HW3_Qty }}",
    
    "HW4_Specs": "{{ HW4_Specs }}",
    "HW4_Maker": "{{ HW4_Maker }}",
    "HW4_Qty": "{{ HW4_Qty }}",
    
    "HW5_Specs": "{{ HW5_Specs }}",
    "HW5_Maker": "{{ HW5_Maker }}",
    "HW5_Qty": "{{ HW5_Qty }}",
    
    "HW6_Specs": "{{ HW6_Specs }}",
    "HW6_Maker": "{{ HW6_Maker }}",
    "HW6_Qty": "{{ HW6_Qty }}",
    
    # Software Specs (9 rows)
    "SW1_Name": "{{ SW1_Name }}",
    "SW1_Maker": "{{ SW1_Maker }}",
    "SW2_Name": "{{ SW2_Name }}",
    "SW2_Maker": "{{ SW2_Maker }}",
    "SW3_Name": "{{ SW3_Name }}",
    "SW3_Maker": "{{ SW3_Maker }}",
    "SW4_Name": "{{ SW4_Name }}",
    "SW4_Maker": "{{ SW4_Maker }}",
    "SW5_Name": "{{ SW5_Name }}",
    "SW5_Maker": "{{ SW5_Maker }}",
    "SW6_Name": "{{ SW6_Name }}",
    "SW6_Maker": "{{ SW6_Maker }}",
    "SW7_Name": "{{ SW7_Name }}",
    "SW7_Maker": "{{ SW7_Maker }}",
    "SW8_Name": "{{ SW8_Name }}",
    "SW8_Maker": "{{ SW8_Maker }}",
    "SW9_Name": "{{ SW9_Name }}",
    "SW9_Maker": "{{ SW9_Maker }}",
    
    # Third Party Software
    "ThirdPartySW": "{{ ThirdPartySW }}",
    
    # Supervisors / Man Days
    "PMDays": "{{ PMDays }}",
    "DevDays": "{{ DevDays }}",
    "CommDays": "{{ CommDays }}",
    "TotalManDays": "{{ TotalManDays }}",
    
    # Value Addition
    "ValueAddedOfferings": "{{ ValueAddedOfferings }}",
    
    # POC
    "POCName": "{{ POCName }}",
    "POCDescription": "{{ POCDescription }}",
}


def replace_in_paragraph(paragraph):
    """
    Replace text in a paragraph while preserving formatting.
    
    This function concatenates all run texts, applies replacements,
    then updates the first run with the new text and clears remaining runs.
    This preserves the paragraph structure while allowing text replacement.
    """
    # Concatenate all run texts
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # Apply all replacements
    modified_text = full_text
    for old_text, new_text in REPLACEMENTS.items():
        modified_text = modified_text.replace(old_text, new_text)
    
    # If text was modified, update the paragraph
    if modified_text != full_text:
        # Clear all runs except the first
        for i in range(len(paragraph.runs) - 1, 0, -1):
            paragraph.runs[i].text = ''
        
        # Update first run with modified text
        if paragraph.runs:
            paragraph.runs[0].text = modified_text
        else:
            # If no runs exist, add one
            paragraph.add_run(modified_text)


def process_document(doc):
    """
    Process all paragraphs in the document body and tables.
    
    This function iterates through:
    1. All paragraphs in the document body
    2. All paragraphs in all table cells
    
    And applies text replacements to each paragraph.
    """
    # Process all paragraphs in document body
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Process all paragraphs in all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def main():
    """
    Main conversion function.
    
    Loads the original template, applies replacements, and saves
    the Jinja2-compatible template.
    """
    # Check if input file exists
    if not INPUT.exists():
        print(f"ERROR: Input file not found: {INPUT}")
        print(f"Please ensure {INPUT} exists before running this script.")
        return
    
    # Create templates directory if it doesn't exist
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    
    # Load the original document
    print(f"Loading template from: {INPUT}")
    doc = Document(INPUT)
    
    # Process the document
    print("Processing document and applying replacements...")
    process_document(doc)
    
    # Save the converted document
    print(f"Saving Jinja2 template to: {OUTPUT}")
    doc.save(OUTPUT)
    
    print("\n" + "="*70)
    print("✓ Template conversion complete!")
    print(f"✓ Output saved to: {OUTPUT}")
    print("="*70)
    
    print("\n⚠️  IMPORTANT: Manual Verification Required")
    print("-" * 70)
    print("1. Open the generated template in Microsoft Word")
    print("2. Navigate to the Features table section")
    print("3. Verify that dynamic feature rows use the following Jinja2 syntax:")
    print("   {%tr for feature in features %}")
    print("   {{ feature.title }}")
    print("   {{ feature.brief }}")
    print("   {{ feature.description }}")
    print("   {%tr endfor %}")
    print("\n4. If the syntax is incorrect, manually add the loop tags")
    print("5. Save the template after verification")
    print("-" * 70)
    
    print("\n📚 Reference Documentation:")
    print("   docxtpl: https://docxtpl.readthedocs.io/en/latest/")
    print("   Jinja2: https://jinja.palletsprojects.com/")
    print("\n")


if __name__ == "__main__":
    main()
