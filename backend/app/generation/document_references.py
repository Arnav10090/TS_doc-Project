"""Figure and table numbering helpers for preview/export parity."""
from __future__ import annotations

import base64
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from html.parser import HTMLParser

from app.generation.completion import strip_html


EDIT_METADATA_KEY = "__editMetadata"


@dataclass(frozen=True)
class DocumentReference:
    id: str
    number: int
    name: str


@dataclass(frozen=True)
class DocumentReferences:
    figures: List[DocumentReference]
    tables: List[DocumentReference]

    @property
    def figure_by_id(self) -> Dict[str, DocumentReference]:
        return {figure.id: figure for figure in self.figures}

    @property
    def table_by_id(self) -> Dict[str, DocumentReference]:
        return {table.id: table for table in self.tables}


PREDEFINED_SECTION_ORDER = [
    "cover",
    "revision_history",
    "executive_summary",
    "introduction",
    "abbreviations",
    "process_flow",
    "overview",
    "features",
    "remote_support",
    "documentation_control",
    "customer_training",
    "system_config",
    "fat_condition",
    "tech_stack",
    "hardware_specs",
    "software_specs",
    "third_party_sw",
    "overall_gantt",
    "shutdown_gantt",
    "supervisors",
    "scope_definitions",
    "division_of_eng",
    "value_addition",
    "work_completion",
    "buyer_obligations",
    "exclusion_list",
    "buyer_prerequisites",
    "binding_conditions",
    "cybersecurity",
    "disclaimer",
    "poc",
]

BUILT_IN_TABLES = {
    "revision_history": ("table:revision_history", "Revision History", 1),
    "abbreviations": ("table:abbreviations", "Abbreviations Used", 3),
    "tech_stack": ("table:tech_stack", "Technology Stack", 4),
    "hardware_specs": ("table:hardware_specs", "Hardware Specifications", 5),
    "software_specs": ("table:software_specs", "Software Specifications", 6),
    "division_of_eng": ("table:division_of_eng", "Responsibility Matrix", 7),
}

BUILT_IN_FIGURES = {
    "system_config": (
        "figure:system_config:architecture",
        "System Architecture",
        "architecture.png",
    ),
    "overall_gantt": (
        "figure:overall_gantt",
        "Overall Gantt Chart",
        "gantt_overall.png",
    ),
    "shutdown_gantt": (
        "figure:shutdown_gantt",
        "Shutdown Gantt Chart",
        "gantt_shutdown.png",
    ),
}


def is_custom_section_key(key: str) -> bool:
    pattern = r"^custom_section_\d+_[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}$"
    return bool(re.match(pattern, key))


def strip_edit_metadata(content: Any) -> Dict[str, Any]:
    if not isinstance(content, dict):
        return {}

    return {key: value for key, value in content.items() if key != EDIT_METADATA_KEY}


def caption_or_fallback(value: Any, fallback: str) -> str:
    caption = str(value).strip() if value is not None else ""
    return caption or fallback


def strip_extension(filename: str) -> str:
    stem = re.sub(r"\.[^.]+$", "", filename or "")
    return re.sub(r"[_-]+", " ", stem).strip()


def get_table_items(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    if isinstance(data.get("tables"), list):
        return [table for table in data["tables"] if isinstance(table, dict)]

    if isinstance(data.get("columns"), list) and isinstance(data.get("rows"), list):
        return [{"columns": data["columns"], "rows": data["rows"]}]

    return []


def get_image_items(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    if isinstance(data.get("images"), list):
        return [image for image in data["images"] if isinstance(image, dict)]

    if data.get("base64"):
        return [
            {
                "base64": data.get("base64", ""),
                "filename": data.get("filename", ""),
                "mimeType": data.get("mimeType", ""),
                "caption": data.get("caption", ""),
            }
        ]

    return []


def _is_inline_custom_section(content: Dict[str, Any]) -> bool:
    return content.get("displayMode") == "subsection"


def _ordered_custom_sections_after(
    custom_sections: Dict[str, Dict[str, Any]],
    insert_after_key: str,
    visited_custom: Set[str],
) -> List[Tuple[str, Dict[str, Any]]]:
    matching = [
        (key, content)
        for key, content in custom_sections.items()
        if key not in visited_custom and content.get("insertAfterKey") == insert_after_key
    ]

    return [
        *[(key, content) for key, content in matching if _is_inline_custom_section(content)],
        *[(key, content) for key, content in matching if not _is_inline_custom_section(content)],
    ]


def collect_document_references(
    all_sections: Dict[str, Any],
    upload_dir: str,
    project_id: str,
) -> DocumentReferences:
    """Collect document-wide references in the same order they appear."""
    clean_sections = {
        key: strip_edit_metadata(content)
        for key, content in all_sections.items()
    }
    custom_sections = {
        key: content
        for key, content in clean_sections.items()
        if is_custom_section_key(key)
    }
    figures: List[DocumentReference] = []
    tables: List[DocumentReference] = []
    visited_custom: Set[str] = set()
    images_dir = Path(upload_dir) / "images" / project_id

    def add_figure(reference_id: str, name: str) -> None:
        figures.append(
            DocumentReference(
                id=reference_id,
                number=len(figures) + 1,
                name=name,
            )
        )

    def add_table(reference_id: str, name: str) -> None:
        tables.append(
            DocumentReference(
                id=reference_id,
                number=len(tables) + 1,
                name=name,
            )
        )

    def collect_custom_after(insert_after_key: str) -> None:
        for section_key, content in _ordered_custom_sections_after(
            custom_sections,
            insert_after_key,
            visited_custom,
        ):
            visited_custom.add(section_key)
            for subsection_index, subsection in enumerate(content.get("subsections", [])):
                if not isinstance(subsection, dict):
                    continue

                subsection_name = str(subsection.get("name") or "").strip()
                content_type = subsection.get("contentType")
                data = subsection.get("data") if isinstance(subsection.get("data"), dict) else {}
                subsection_key = subsection.get("key") or f"subsection-{subsection_index}"

                if content_type == "table":
                    for table_index, table in enumerate(get_table_items(data)):
                        add_table(
                            f"table:{section_key}:{subsection_key}:{table_index}",
                            caption_or_fallback(
                                table.get("caption"),
                                subsection_name
                                or f"Table {subsection_index + 1}{'.' + str(table_index + 1) if table_index else ''}",
                            ),
                        )

                if content_type == "image":
                    for image_index, image in enumerate(get_image_items(data)):
                        if not image.get("base64"):
                            continue

                        add_figure(
                            f"figure:{section_key}:{subsection_key}:{image_index}",
                            caption_or_fallback(
                                image.get("caption"),
                                subsection_name
                                or strip_extension(str(image.get("filename") or ""))
                                or f"Figure {subsection_index + 1}{'.' + str(image_index + 1) if image_index else ''}",
                            ),
                        )

            collect_custom_after(section_key)

    for section_key in PREDEFINED_SECTION_ORDER:
        if section_key not in clean_sections:
            continue

        built_in_table = BUILT_IN_TABLES.get(section_key)
        if built_in_table:
            reference_id, name, _table_index = built_in_table
            add_table(reference_id, name)

        built_in_figure = BUILT_IN_FIGURES.get(section_key)
        if built_in_figure:
            reference_id, name, filename = built_in_figure
            if (images_dir / filename).exists():
                add_figure(reference_id, name)

        collect_custom_after(section_key)

    for section_key, _content in custom_sections.items():
        if section_key not in visited_custom:
            collect_custom_after(custom_sections[section_key].get("insertAfterKey", ""))

    return DocumentReferences(figures=figures, tables=tables)


def _format_caption(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.size = Pt(10)


def _insert_paragraph_before_table(table: Table, text: str) -> Paragraph:
    paragraph = table._parent.add_paragraph(text)
    table._tbl.addprevious(paragraph._p)
    _format_caption(paragraph)
    return paragraph


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    _format_caption(inserted)
    return inserted


def _has_drawing(paragraph: Paragraph) -> bool:
    return next(paragraph._p.iter(qn("w:drawing")), None) is not None


def _find_next_drawing_after_text(
    doc: DocumentObject,
    marker: str,
    max_scan: int = 5,
) -> Optional[Paragraph]:
    marker_lower = marker.lower()
    for index, paragraph in enumerate(doc.paragraphs):
        if marker_lower not in paragraph.text.lower():
            continue

        for candidate in doc.paragraphs[index + 1 : index + max_scan + 1]:
            if _has_drawing(candidate):
                return candidate
        return None

    return None


def _find_previous_drawing_before_text(
    doc: DocumentObject,
    marker: str,
    max_scan: int = 6,
) -> Optional[Paragraph]:
    marker_lower = marker.lower()
    for index, paragraph in enumerate(doc.paragraphs):
        if marker_lower not in paragraph.text.lower():
            continue

        lower_bound = max(0, index - max_scan)
        for candidate in reversed(doc.paragraphs[lower_bound:index]):
            if _has_drawing(candidate):
                return candidate
        return None

    return None


def _set_cant_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _fill_table(table: Table, headers: List[str], rows: Iterable[List[str]]) -> None:
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value

    for row in table.rows:
        _set_cant_split(row)


def _apply_table_grid_style(table: Table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def _add_heading_paragraph(
    doc: DocumentObject,
    text: str,
    font_size: int,
    style_name: Optional[str] = None,
) -> Paragraph:
    paragraph = doc.add_paragraph()
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            pass
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    return paragraph


def _next_heading_number(doc: DocumentObject, style_name: str) -> int:
    return (
        sum(
            1
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
            and paragraph.style is not None
            and paragraph.style.name == style_name
        )
        + 1
    )


def _set_update_fields_on_open(doc: DocumentObject) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))

    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def _insert_toc_field_after(paragraph: Paragraph) -> None:
    field_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(field_paragraph_element)
    field_paragraph = Paragraph(field_paragraph_element, paragraph._parent)

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    field_paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    instruction_run.append(instruction)
    field_paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    field_paragraph._p.append(separate_run)

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Update field to refresh table of contents."
    placeholder_run.append(placeholder_text)
    field_paragraph._p.append(placeholder_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    field_paragraph._p.append(end_run)


def _element_contains_toc_field(element: Any) -> bool:
    for instruction in element.iter(qn("w:instrText")):
        if "TOC" in (instruction.text or ""):
            return True

    for field in element.iter(qn("w:fldSimple")):
        if "TOC" in (field.get(qn("w:instr")) or ""):
            return True

    return False


def _apply_dynamic_table_of_contents(doc: DocumentObject) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().upper() != "TABLE OF CONTENTS":
            continue

        next_element = paragraph._p.getnext()
        if next_element is not None and _element_contains_toc_field(next_element):
            _set_update_fields_on_open(doc)
            return

        _insert_toc_field_after(paragraph)
        _set_update_fields_on_open(doc)
        return


def _add_reference_table(
    doc: DocumentObject,
    headers: List[str],
    references: List[DocumentReference],
    prefix: str,
) -> None:
    table = doc.add_table(rows=1, cols=3)
    _apply_table_grid_style(table)
    _fill_table(
        table,
        headers,
        [
            [str(index + 1), f"{prefix} {reference.number}", reference.name]
            for index, reference in enumerate(references)
        ],
    )


def _decode_image_to_file(
    image: Dict[str, Any],
    upload_dir: str,
    project_id: str,
    image_index: int,
) -> Optional[Path]:
    base64_data = str(image.get("base64") or "")
    if not base64_data:
        return None

    if "," in base64_data:
        base64_data = base64_data.split(",", 1)[1]

    try:
        image_bytes = base64.b64decode(base64_data)
    except Exception:
        return None

    filename = str(image.get("filename") or f"custom-image-{image_index}.png")
    safe_filename = re.sub(r"[^A-Za-z0-9_.-]", "_", filename)
    image_dir = Path(upload_dir) / "images" / project_id / "custom"
    image_dir.mkdir(parents=True, exist_ok=True)
    image_path = image_dir / safe_filename
    image_path.write_bytes(image_bytes)
    return image_path


SECTION_HEADING_MARKERS: Dict[str, List[str]] = {
    "revision_history": ["REVISION HISTORY"],
    "executive_summary": ["EXECUTIVE SUMMARY"],
    "introduction": ["INTRODUCTION"],
    "abbreviations": ["ABBREVIATIONS USED", "ABBREVIATIONS"],
    "process_flow": ["PROCESS FLOW"],
    "overview": ["OVERVIEW OF"],
    "features": ["DESIGN SCOPE OF WORK"],
    "remote_support": ["REMOTE SUPPORT SYSTEM"],
    "documentation_control": ["DOCUMENTATION CONTROL"],
    "customer_training": ["CUSTOMER TRAINING"],
    "system_config": ["SYSTEM CONFIGURATION"],
    "fat_condition": ["FAT CONDITION"],
    "tech_stack": ["TECHNOLOGY STACK"],
    "hardware_specs": ["BASIC HARDWARE SPECIFICATIONS", "HARDWARE SPECIFICATIONS"],
    "software_specs": ["BASIC SOFTWARE SPECIFICATION", "SOFTWARE SPECIFICATIONS"],
    "third_party_sw": ["THIRD PARTY SOFTWARE"],
    "overall_gantt": ["OVERALL GANTT"],
    "shutdown_gantt": ["SHUTDOWN GANTT"],
    "supervisors": ["SUPERVISORS"],
    "scope_definitions": ["SCOPE OF SUPPLY DEFINITIONS"],
    "division_of_eng": ["DIVISION OF ENGINEERING"],
    "value_addition": ["VALUE ADDITION"],
    "work_completion": ["WORK COMPLETION"],
    "buyer_obligations": ["BUYER OBLIGATIONS"],
    "exclusion_list": ["EXCLUSION LIST"],
    "buyer_prerequisites": ["BUYER PREREQUISITES"],
    "binding_conditions": ["BINDING CONDITIONS"],
    "cybersecurity": ["CYBERSECURITY"],
    "disclaimer": ["DISCLAIMER"],
    "poc": ["PROOF OF CONCEPT", "PoC"],
}


class _RichTextBlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: List[Tuple[str, str]] = []
        self._list_stack: List[str] = []
        self._current_kind: Optional[str] = None
        self._buffer: List[str] = []
        self._in_li = False

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        if tag in {"ul", "ol"}:
            self._list_stack.append(tag)
            return

        if tag == "li":
            self._start_block(self._list_stack[-1] if self._list_stack else "ul")
            self._in_li = True
            return

        if tag == "p" and not self._in_li:
            self._start_block("p")
            return

        if tag == "br" and self._current_kind:
            self._buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"ul", "ol"}:
            if self._list_stack:
                self._list_stack.pop()
            return

        if tag == "li":
            self._finish_block()
            self._in_li = False
            return

        if tag == "p" and self._current_kind == "p" and not self._in_li:
            self._finish_block()

    def handle_data(self, data: str) -> None:
        if not self._current_kind and data.strip():
            self._start_block("p")
        if self._current_kind:
            self._buffer.append(data)

    def close(self) -> None:
        super().close()
        self._finish_block()

    def _start_block(self, kind: str) -> None:
        self._finish_block()
        self._current_kind = kind
        self._buffer = []

    def _finish_block(self) -> None:
        if not self._current_kind:
            return

        text = re.sub(r"\s+", " ", "".join(self._buffer)).strip()
        if text:
            self.blocks.append((self._current_kind, text))
        self._current_kind = None
        self._buffer = []


def _html_blocks(html: str) -> List[Tuple[str, str]]:
    parser = _RichTextBlockParser()
    parser.feed(html or "")
    parser.close()
    if parser.blocks:
        return parser.blocks

    text = strip_html(html or "").strip()
    return [("p", text)] if text else []


def _paragraph_heading_level(paragraph: Paragraph) -> Optional[int]:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    match = re.match(r"Heading\s+(\d+)", style_name)
    return int(match.group(1)) if match else None


def _paragraph_element_heading_level(
    doc: DocumentObject,
    element: Any,
) -> Optional[int]:
    if element.tag != qn("w:p"):
        return None

    return _paragraph_heading_level(Paragraph(element, doc))


def _matches_marker(text: str, marker: str) -> bool:
    normalized_text = re.sub(r"\s+", " ", text).strip().lower()
    normalized_marker = re.sub(r"\s+", " ", marker).strip().lower()
    return normalized_marker in normalized_text


def _section_markers(section_key: str, all_sections: Dict[str, Any]) -> List[str]:
    content = strip_edit_metadata(all_sections.get(section_key, {}))
    markers = []
    heading = content.get("heading")
    if isinstance(heading, str) and heading.strip():
        markers.append(heading)

    markers.extend(SECTION_HEADING_MARKERS.get(section_key, []))
    return markers


def _last_element_in_section(doc: DocumentObject, heading: Paragraph) -> Any:
    heading_level = _paragraph_heading_level(heading) or 1
    cursor = heading._p
    next_element = cursor.getnext()

    while next_element is not None and next_element.tag != qn("w:sectPr"):
        next_heading_level = _paragraph_element_heading_level(doc, next_element)
        if next_heading_level is not None and next_heading_level <= heading_level:
            break

        cursor = next_element
        next_element = next_element.getnext()

    return cursor


def _find_section_insert_cursor(
    doc: DocumentObject,
    section_key: str,
    all_sections: Dict[str, Any],
) -> Optional[Any]:
    markers = _section_markers(section_key, all_sections)
    if not markers:
        return None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if any(_matches_marker(text, marker) for marker in markers):
            return _last_element_in_section(doc, paragraph)

    return None


def _set_paragraph_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        pass


def _style_exists(doc: DocumentObject, style_name: str) -> bool:
    try:
        doc.styles[style_name]
    except KeyError:
        return False

    return True


def _insert_paragraph_after_element(
    doc: DocumentObject,
    cursor: Any,
    text: str,
    style_name: Optional[str] = None,
) -> Tuple[Paragraph, Any]:
    paragraph = doc.add_paragraph()
    if style_name:
        _set_paragraph_style(paragraph, style_name)
    paragraph.add_run(text)
    cursor.addnext(paragraph._p)
    return paragraph, paragraph._p


def _insert_table_after_element(
    doc: DocumentObject,
    cursor: Any,
    columns: List[str],
    rows: Iterable[List[str]],
) -> Any:
    table = doc.add_table(rows=1, cols=len(columns))
    _apply_table_grid_style(table)
    _fill_table(table, columns, rows)
    cursor.addnext(table._tbl)
    return table._tbl


def _insert_picture_after_element(
    doc: DocumentObject,
    cursor: Any,
    image_path: Path,
) -> Any:
    doc.add_picture(str(image_path), width=Cm(15))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.addnext(paragraph._p)
    return paragraph._p


def _insert_rich_text_after_element(
    doc: DocumentObject,
    cursor: Any,
    html: str,
) -> Any:
    current = cursor
    previous_kind: Optional[str] = None
    ordered_index = 0

    for kind, text in _html_blocks(html):
        if kind == "ol":
            ordered_index = ordered_index + 1 if previous_kind == "ol" else 1
        else:
            ordered_index = 0

        style_name = {
            "ul": "List Bullet",
            "ol": "List Number",
        }.get(kind, "Normal")

        if kind == "ul" and not _style_exists(doc, style_name):
            text = f"• {text}"
            style_name = "Normal"
        elif kind == "ol" and not _style_exists(doc, style_name):
            text = f"{ordered_index}. {text}"
            style_name = "Normal"

        _paragraph, current = _insert_paragraph_after_element(
            doc,
            current,
            text,
            style_name,
        )
        previous_kind = kind

    return current


def _custom_section_title(content: Dict[str, Any]) -> str:
    title = str(content.get("title") or "").strip()
    if title and title.upper() != "NEW SECTION":
        return title

    for subsection in content.get("subsections", []):
        if isinstance(subsection, dict):
            name = str(subsection.get("name") or "").strip()
            if name:
                return name

    return "New Section"


def _insert_custom_subsection(
    doc: DocumentObject,
    cursor: Any,
    section_key: str,
    subsection: Dict[str, Any],
    subsection_index: int,
    references: DocumentReferences,
    upload_dir: str,
    project_id: str,
) -> Any:
    table_by_id = references.table_by_id
    figure_by_id = references.figure_by_id
    subsection_key = subsection.get("key") or f"subsection-{subsection_index}"
    data = subsection.get("data") if isinstance(subsection.get("data"), dict) else {}
    heading = caption_or_fallback(subsection.get("name"), "New Subsection")
    _heading_paragraph, current = _insert_paragraph_after_element(
        doc,
        cursor,
        heading,
        "Heading 2",
    )

    if subsection.get("contentType") == "table":
        for table_index, table_data in enumerate(get_table_items(data)):
            reference = table_by_id.get(f"table:{section_key}:{subsection_key}:{table_index}")
            if reference:
                caption, current = _insert_paragraph_after_element(
                    doc,
                    current,
                    f"Table {reference.number}: {reference.name}",
                )
                _format_caption(caption)

            columns = [str(column) for column in table_data.get("columns", [])]
            rows = table_data.get("rows", [])
            if columns:
                current = _insert_table_after_element(
                    doc,
                    current,
                    columns,
                    [
                        [str(row.get(column, "")) for column in columns]
                        for row in rows
                        if isinstance(row, dict)
                    ],
                )

    if subsection.get("contentType") == "image":
        for image_index, image in enumerate(get_image_items(data)):
            image_path = _decode_image_to_file(
                image,
                upload_dir,
                project_id,
                image_index,
            )
            if image_path:
                current = _insert_picture_after_element(doc, current, image_path)

            reference = figure_by_id.get(
                f"figure:{section_key}:{subsection_key}:{image_index}"
            )
            if reference:
                caption, current = _insert_paragraph_after_element(
                    doc,
                    current,
                    f"Figure {reference.number}: {reference.name}",
                )
                _format_caption(caption)

    if subsection.get("contentType") == "paragraph":
        paragraphs = data.get("paragraphs")
        if not isinstance(paragraphs, list):
            paragraphs = [{"html": data.get("html", "")}]

        for paragraph_data in paragraphs:
            if isinstance(paragraph_data, dict):
                current = _insert_rich_text_after_element(
                    doc,
                    current,
                    str(paragraph_data.get("html") or ""),
                )

    return current


def _insert_custom_sections(
    doc: DocumentObject,
    all_sections: Dict[str, Any],
    references: DocumentReferences,
    upload_dir: str,
    project_id: str,
) -> None:
    custom_sections = {
        key: strip_edit_metadata(content)
        for key, content in all_sections.items()
        if is_custom_section_key(key)
    }
    visited_custom: Set[str] = set()

    def insert_after(anchor_key: str, cursor: Any) -> Any:
        current = cursor
        for section_key, content in _ordered_custom_sections_after(
            custom_sections,
            anchor_key,
            visited_custom,
        ):
            visited_custom.add(section_key)

            if _is_inline_custom_section(content):
                for subsection_index, subsection in enumerate(content.get("subsections", [])):
                    if isinstance(subsection, dict):
                        current = _insert_custom_subsection(
                            doc,
                            current,
                            section_key,
                            subsection,
                            subsection_index,
                            references,
                            upload_dir,
                            project_id,
                        )
                current = insert_after(section_key, current)
                continue

            _heading_paragraph, current = _insert_paragraph_after_element(
                doc,
                current,
                _custom_section_title(content),
                "Heading 1",
            )
            for subsection_index, subsection in enumerate(content.get("subsections", [])):
                if isinstance(subsection, dict):
                    current = _insert_custom_subsection(
                        doc,
                        current,
                        section_key,
                        subsection,
                        subsection_index,
                        references,
                        upload_dir,
                        project_id,
                    )
            current = insert_after(section_key, current)

        return current

    for section_key in PREDEFINED_SECTION_ORDER:
        if section_key not in all_sections:
            continue

        cursor = _find_section_insert_cursor(doc, section_key, all_sections)
        if cursor is not None:
            insert_after(section_key, cursor)

    fallback_cursor = doc.paragraphs[-1]._p if doc.paragraphs else doc.add_paragraph()._p
    for section_key, content in custom_sections.items():
        if section_key in visited_custom:
            continue

        visited_custom.add(section_key)
        _heading_paragraph, fallback_cursor = _insert_paragraph_after_element(
            doc,
            fallback_cursor,
            _custom_section_title(content),
            "Heading 1",
        )
        for subsection_index, subsection in enumerate(content.get("subsections", [])):
            if isinstance(subsection, dict):
                fallback_cursor = _insert_custom_subsection(
                    doc,
                    fallback_cursor,
                    section_key,
                    subsection,
                    subsection_index,
                    references,
                    upload_dir,
                    project_id,
                )
        fallback_cursor = insert_after(section_key, fallback_cursor)


def apply_document_references(
    file_path: str,
    all_sections: Dict[str, Any],
    upload_dir: str,
    project_id: str,
) -> DocumentReferences:
    """Add captions and the final list section to an exported DOCX file."""
    references = collect_document_references(all_sections, upload_dir, project_id)
    doc = Document(file_path)
    _apply_dynamic_table_of_contents(doc)
    table_by_id = references.table_by_id
    figure_by_id = references.figure_by_id

    for section_key, (reference_id, _name, table_index) in BUILT_IN_TABLES.items():
        if section_key not in all_sections:
            continue

        reference = table_by_id.get(reference_id)
        if reference and table_index < len(doc.tables):
            _insert_paragraph_before_table(
                doc.tables[table_index],
                f"Table {reference.number}: {reference.name}",
            )

    architecture = figure_by_id.get("figure:system_config:architecture")
    if architecture:
        paragraph = _find_next_drawing_after_text(
            doc,
            "reference system configuration",
        )
        if paragraph:
            _insert_paragraph_after(
                paragraph,
                f"Figure {architecture.number}: {architecture.name}",
            )

    overall = figure_by_id.get("figure:overall_gantt")
    if overall:
        paragraph = _find_previous_drawing_before_text(
            doc,
            "After Approval on System Design Document",
        )
        if paragraph:
            _insert_paragraph_after(
                paragraph,
                f"Figure {overall.number}: {overall.name}",
            )

    shutdown = figure_by_id.get("figure:shutdown_gantt")
    if shutdown:
        paragraph = _find_previous_drawing_before_text(doc, "NOTE:")
        if paragraph:
            _insert_paragraph_after(
                paragraph,
                f"Figure {shutdown.number}: {shutdown.name}",
            )

    _insert_custom_sections(doc, all_sections, references, upload_dir, project_id)

    doc.add_page_break()
    reference_section_number = _next_heading_number(doc, "Heading 1")
    _add_heading_paragraph(
        doc,
        f"{reference_section_number}. List of Figures and Tables",
        16,
        "Heading 1",
    )
    _add_heading_paragraph(
        doc,
        f"{reference_section_number}.1 List of Figures",
        12,
        "Heading 2",
    )
    _add_reference_table(
        doc,
        ["S No.", "Figure No.", "Figure Name"],
        references.figures,
        "Figure",
    )
    _add_heading_paragraph(
        doc,
        f"{reference_section_number}.2 List of Tables",
        12,
        "Heading 2",
    )
    _add_reference_table(
        doc,
        ["S No.", "Table No.", "Table Name"],
        references.tables,
        "Table",
    )

    doc.save(file_path)
    return references
